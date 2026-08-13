import ast
import re
from abc import ABC, abstractmethod
from pathlib import Path

from app.services.ingestion.models import DocumentElement, ParsedDocument


CODE_LANGUAGES = {
    ".py": "python", ".js": "javascript", ".jsx": "javascript",
    ".ts": "typescript", ".tsx": "typescript", ".java": "java",
    ".go": "go", ".rs": "rust", ".c": "c", ".h": "c",
    ".cpp": "cpp", ".hpp": "cpp", ".cs": "csharp", ".php": "php",
    ".vue": "vue", ".svelte": "svelte", ".sql": "sql",
    ".html": "html", ".css": "css", ".scss": "scss",
    ".json": "json", ".yaml": "yaml", ".yml": "yaml", ".toml": "toml",
}


def decode_text(data: bytes) -> tuple[str, str]:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return data.decode(encoding), encoding
        except UnicodeDecodeError:
            continue
    raise ValueError("无法识别文件编码，请转换为 UTF-8 后重试")


class BaseParser(ABC):
    name = "base"
    content_type = "text"

    @abstractmethod
    def parse(self, path: Path) -> ParsedDocument:
        raise NotImplementedError


class TextParser(BaseParser):
    name = "plain-text-v1"

    def parse(self, path: Path) -> ParsedDocument:
        text, encoding = decode_text(path.read_bytes())
        element_type = "markdown" if path.suffix.lower() == ".md" else "paragraph"
        return ParsedDocument(
            content_type="text",
            parser_name=self.name,
            elements=[DocumentElement(element_type, text)],
            metadata={"encoding": encoding},
        )


class CodeParser(BaseParser):
    name = "source-code-v1"
    content_type = "code"

    def parse(self, path: Path) -> ParsedDocument:
        source, encoding = decode_text(path.read_bytes())
        language = CODE_LANGUAGES[path.suffix.lower()]
        elements = self._python_elements(source) if language == "python" else []
        if not elements:
            elements = [DocumentElement("source", source, {
                "language": language, "start_line": 1,
                "end_line": len(source.splitlines()),
            })]
        return ParsedDocument("code", self.name, elements, {
            "encoding": encoding, "language": language,
        })

    @staticmethod
    def _python_elements(source: str) -> list[DocumentElement]:
        try:
            tree = ast.parse(source)
        except SyntaxError:
            return []
        lines = source.splitlines()
        elements = []
        definitions = [node for node in tree.body if isinstance(
            node, (ast.FunctionDef, ast.AsyncFunctionDef, ast.ClassDef)
        )]
        first_line = min((node.lineno for node in definitions), default=len(lines) + 1)
        if first_line > 1:
            preamble = "\n".join(lines[:first_line - 1]).strip()
            if preamble:
                elements.append(DocumentElement("preamble", preamble, {
                    "language": "python", "start_line": 1, "end_line": first_line - 1,
                }))
        for node in definitions:
            start = node.lineno
            end = getattr(node, "end_lineno", start)
            elements.append(DocumentElement(
                "class" if isinstance(node, ast.ClassDef) else "function",
                "\n".join(lines[start - 1:end]),
                {"language": "python", "symbol_name": node.name,
                 "start_line": start, "end_line": end},
            ))
        return elements


class WordParser(BaseParser):
    name = "python-docx-v1"
    content_type = "document"

    def parse(self, path: Path) -> ParsedDocument:
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("缺少 python-docx 依赖，无法解析 Word 文档") from exc
        document = Document(path)
        elements = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style = paragraph.style.name if paragraph.style else ""
            match = re.match(r"(?:Heading|标题)\s*(\d+)", style, re.IGNORECASE)
            if match:
                elements.append(DocumentElement("heading", text, {"level": int(match.group(1))}))
            else:
                elements.append(DocumentElement("paragraph", text))
        for index, table in enumerate(document.tables, 1):
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            rendered = "\n".join(" | ".join(row) for row in rows if any(row))
            if rendered:
                elements.append(DocumentElement("table", rendered, {"table_index": index}))
        return ParsedDocument("document", self.name, elements)


class PdfParser(BaseParser):
    name = "pymupdf-v1"
    content_type = "document"

    def parse(self, path: Path) -> ParsedDocument:
        try:
            import pymupdf as fitz
        except ImportError as exc:
            raise RuntimeError("缺少 PyMuPDF 依赖，无法解析 PDF") from exc
        elements = []
        with fitz.open(path) as document:
            for page_number, page in enumerate(document, 1):
                blocks = sorted(page.get_text("blocks"), key=lambda block: (block[1], block[0]))
                for block in blocks:
                    text = block[4].strip()
                    if text:
                        elements.append(DocumentElement("paragraph", text, {
                            "page": page_number,
                            "bbox": [round(value, 2) for value in block[:4]],
                        }))
        if not elements:
            raise ValueError("PDF 中未提取到文本，可能是扫描件；当前版本暂未启用 OCR")
        return ParsedDocument("document", self.name, elements)
